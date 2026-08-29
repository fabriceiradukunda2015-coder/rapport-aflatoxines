import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
import datetime

st.set_page_config(page_title="Validation Aflatoxines B2", layout="wide")
st.title("🧪 Rapport de Validation - Aflatoxines B2")
st.write("Charge ton fichier Excel pour générer le rapport automatiquement")

fichier = st.file_uploader("📁 Choisir fichier Excel", type=['xlsx'])

def generer_word(r2, ld, lq, recov, rsd_rep, rsd_pi):
    doc = Document()
    doc.add_heading('RAPPORT DE VALIDATION DE METHODE', 0)
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"R²: {r2} - {'Conforme' if r2 >= 0.98 else 'Non Conforme'}")
    doc.add_paragraph(f"LD: {ld} ppb | LQ: {lq} ppb")
    doc.add_paragraph(f"Recouvrement: {recov}%")
    doc.add_paragraph(f"Répétabilité RSD: {rsd_rep}%")
    doc.add_paragraph(f"Précision Intermédiaire RSD: {rsd_pi}%")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

if fichier:
    try:
        xls = pd.ExcelFile(fichier)
        
        df_lin = pd.read_excel(xls, sheet_name="Linéarité", header=None)
        df_ld = pd.read_excel(xls, sheet_name="LD-LQ", header=None)
        df_ex = pd.read_excel(xls, sheet_name="Exactitude", header=None)
        df_rep = pd.read_excel(xls, sheet_name="Répétabilité", header=None)
        df_pi = pd.read_excel(xls, sheet_name="Précision intermédiaire", header=None)

        r2 = df_lin.iloc[5, 1]
        ld = df_ld.iloc[0, 1]
        lq = df_ld.iloc[1, 1]
        recov = df_ex.iloc[3, 2]
        rsd_rep = df_rep.iloc[8, 1]
        rsd_pi = df_pi.iloc[6, 1]

        st.success("✅ Fichier analysé avec succès !")
        
        col1, col2, col3 = st.columns(3)
        with col1: st.metric("R²", r2, "≥ 0.98")
        with col2: st.metric("LQ", f"{lq} ppb")
        with col3: st.metric("Recouvrement", f"{recov}%")
        
        col4, col5 = st.columns(2)
        with col4: st.metric("Répétabilité", f"{rsd_rep}%", "≤ 30%")
        with col5: st.metric("Précision Inter", f"{rsd_pi}%", "≤ 45%")

        word_data = generer_word(r2, ld, lq, recov, rsd_rep, rsd_pi)
        st.download_button(
            "📥 Télécharger Rapport Word",
            data=word_data,
            file_name=f"Rapport_Aflatoxines_{datetime.date.today()}.docx"
        )

    except Exception as e:
        st.error(f"Erreur: {e}. Vérifie les noms d'onglets dans Excel")
else:
    st.info("👆 Envoie d'abord ton fichier Excel de validation")
